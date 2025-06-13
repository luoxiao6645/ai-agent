"""
文档解析工具
"""
import asyncio
import logging
import os
from typing import Optional, Type
import io

from langchain.tools.base import BaseTool
from pydantic import BaseModel, Field

# 导入文档处理库
try:
    import PyPDF2
    from docx import Document
    import openpyxl
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("Document processing libraries not available")

logger = logging.getLogger(__name__)

class DocumentParserInput(BaseModel):
    """文档解析输入模型"""
    file_path: str = Field(description="文档文件路径")
    parse_type: str = Field(default="auto", description="解析类型: auto, pdf, docx, xlsx, txt")

class DocumentParserTool(BaseTool):
    """文档解析工具 - 支持PDF、Word、Excel等文档处理"""
    
    name = "document_parser"
    description = "解析和提取文档内容，支持PDF、Word、Excel、文本文件"
    args_schema: Type[BaseModel] = DocumentParserInput
    
    def __init__(self):
        super().__init__()
        self.supported_formats = ['.pdf', '.docx', '.xlsx', '.txt', '.md']
    
    def _run(self, file_path: str, parse_type: str = "auto") -> str:
        """同步执行文档解析"""
        return asyncio.run(self._arun(file_path, parse_type))
    
    async def _arun(self, file_path: str, parse_type: str = "auto") -> str:
        """异步执行文档解析"""
        try:
            # 检查文件是否存在
            if not os.path.exists(file_path):
                return f"文件不存在: {file_path}"
            
            # 检查文件大小
            file_size = os.path.getsize(file_path)
            if file_size > 100 * 1024 * 1024:  # 100MB限制
                return f"文件过大: {file_size / 1024 / 1024:.1f}MB，超过100MB限制"
            
            # 自动检测文件类型
            if parse_type == "auto":
                parse_type = self._detect_file_type(file_path)
            
            # 根据文件类型解析
            if parse_type == "pdf":
                return await self._parse_pdf(file_path)
            elif parse_type == "docx":
                return await self._parse_docx(file_path)
            elif parse_type == "xlsx":
                return await self._parse_xlsx(file_path)
            elif parse_type == "txt":
                return await self._parse_text(file_path)
            else:
                return f"不支持的文件类型: {parse_type}"
                
        except Exception as e:
            logger.error(f"Document parsing failed: {e}")
            return f"文档解析失败: {str(e)}"
    
    def _detect_file_type(self, file_path: str) -> str:
        """检测文件类型"""
        _, ext = os.path.splitext(file_path.lower())
        
        type_mapping = {
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.doc': 'docx',
            '.xlsx': 'xlsx',
            '.xls': 'xlsx',
            '.txt': 'txt',
            '.md': 'txt'
        }
        
        return type_mapping.get(ext, 'txt')
    
    async def _parse_pdf(self, file_path: str) -> str:
        """解析PDF文件"""
        if not PDF_AVAILABLE:
            return "PDF解析库未安装，请安装PyPDF2"
        
        try:
            content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # 获取文档信息
                num_pages = len(pdf_reader.pages)
                content.append(f"PDF文档信息:")
                content.append(f"- 页数: {num_pages}")
                
                if pdf_reader.metadata:
                    if pdf_reader.metadata.title:
                        content.append(f"- 标题: {pdf_reader.metadata.title}")
                    if pdf_reader.metadata.author:
                        content.append(f"- 作者: {pdf_reader.metadata.author}")
                
                content.append("\n文档内容:")
                
                # 提取文本内容
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            content.append(f"\n--- 第{page_num + 1}页 ---")
                            content.append(page_text.strip())
                    except Exception as e:
                        content.append(f"\n--- 第{page_num + 1}页解析失败: {str(e)} ---")
            
            return "\n".join(content)
            
        except Exception as e:
            return f"PDF解析失败: {str(e)}"
    
    async def _parse_docx(self, file_path: str) -> str:
        """解析Word文档"""
        if not PDF_AVAILABLE:
            return "Word解析库未安装，请安装python-docx"
        
        try:
            doc = Document(file_path)
            content = []
            
            content.append("Word文档内容:")
            
            # 提取段落内容
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    content.append(f"\n段落{i + 1}: {paragraph.text}")
            
            # 提取表格内容
            if doc.tables:
                content.append("\n\n表格内容:")
                for table_num, table in enumerate(doc.tables):
                    content.append(f"\n--- 表格{table_num + 1} ---")
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        content.append(" | ".join(row_data))
            
            return "\n".join(content)
            
        except Exception as e:
            return f"Word文档解析失败: {str(e)}"
    
    async def _parse_xlsx(self, file_path: str) -> str:
        """解析Excel文件"""
        if not PDF_AVAILABLE:
            return "Excel解析库未安装，请安装openpyxl"
        
        try:
            workbook = openpyxl.load_workbook(file_path)
            content = []
            
            content.append("Excel文档内容:")
            content.append(f"工作表数量: {len(workbook.sheetnames)}")
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                content.append(f"\n--- 工作表: {sheet_name} ---")
                
                # 获取有数据的区域
                max_row = sheet.max_row
                max_col = sheet.max_column
                
                if max_row > 0 and max_col > 0:
                    content.append(f"数据范围: {max_row}行 x {max_col}列")
                    
                    # 提取前10行数据作为示例
                    for row in sheet.iter_rows(min_row=1, max_row=min(10, max_row), values_only=True):
                        row_data = [str(cell) if cell is not None else "" for cell in row]
                        content.append(" | ".join(row_data))
                    
                    if max_row > 10:
                        content.append(f"... (还有{max_row - 10}行数据)")
                else:
                    content.append("工作表为空")
            
            return "\n".join(content)
            
        except Exception as e:
            return f"Excel文档解析失败: {str(e)}"
    
    async def _parse_text(self, file_path: str) -> str:
        """解析文本文件"""
        try:
            # 尝试不同的编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        
                        # 添加文件信息
                        file_info = f"文本文件信息:\n"
                        file_info += f"- 文件大小: {len(content)}字符\n"
                        file_info += f"- 行数: {content.count(chr(10)) + 1}\n"
                        file_info += f"- 编码: {encoding}\n\n"
                        file_info += "文件内容:\n"
                        file_info += "-" * 50 + "\n"
                        
                        return file_info + content
                        
                except UnicodeDecodeError:
                    continue
            
            return "无法解析文本文件，可能是编码问题"
            
        except Exception as e:
            return f"文本文件解析失败: {str(e)}"
    
    def get_supported_formats(self) -> list:
        """获取支持的文件格式"""
        return self.supported_formats.copy()
    
    async def validate_file(self, file_path: str) -> dict:
        """验证文件"""
        result = {
            "valid": False,
            "file_exists": False,
            "file_size": 0,
            "file_type": None,
            "supported": False,
            "error": None
        }
        
        try:
            # 检查文件是否存在
            if os.path.exists(file_path):
                result["file_exists"] = True
                result["file_size"] = os.path.getsize(file_path)
                
                # 检测文件类型
                file_type = self._detect_file_type(file_path)
                result["file_type"] = file_type
                
                # 检查是否支持
                _, ext = os.path.splitext(file_path.lower())
                result["supported"] = ext in self.supported_formats
                
                # 检查文件大小
                if result["file_size"] <= 100 * 1024 * 1024:
                    result["valid"] = result["supported"]
                else:
                    result["error"] = "文件过大"
            else:
                result["error"] = "文件不存在"
                
        except Exception as e:
            result["error"] = str(e)
        
        return result
